import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn

def create_comprehensive_report():
    doc = Document()
    
    # helper for centering
    def center(p):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    # 1. TITLE PAGE
    title_p = center(doc.add_paragraph())
    run = title_p.add_run("A CAPSTONE PROJECT REPORT ON\n\n")
    run.font.size = Pt(14)
    run.font.bold = True
    
    run = title_p.add_run("DOC 2 ME: HEALTH REPORT SIMPLIFIER USING MULTIMODAL LARGE LANGUAGE MODELS\n\n\n")
    run.font.size = Pt(20)
    run.font.bold = True
    
    run = title_p.add_run("Submitted in partial fulfilment of the requirements for the award of the degree of\n")
    run.font.size = Pt(12)
    run = title_p.add_run("BACHELOR OF TECHNOLOGY\n")
    run.font.size = Pt(14)
    run.font.bold = True
    run = title_p.add_run("IN\nCOMPUTER SCIENCE AND ENGINEERING\n\n\n")
    run.font.size = Pt(14)
    run.font.bold = True

    run = title_p.add_run("Submitted by:\n")
    run.font.size = Pt(12)
    run = title_p.add_run("KUSHAL DEVRAJ (BU22CSEN0102284)\n\n\n")
    run.font.size = Pt(14)
    run.font.bold = True

    run = title_p.add_run("Under the Guidance of:\n")
    run.font.size = Pt(12)
    run = title_p.add_run("SANGAMESHWAR\nAssistant Professor\n\n\n")
    run.font.size = Pt(14)
    run.font.bold = True

    run = title_p.add_run("DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING\nGITAM SCHOOL OF TECHNOLOGY\nGITAM (DEEMED TO BE UNIVERSITY)\nBENGALURU CAMPUS\nAPRIL 2026")
    run.font.size = Pt(12)
    
    doc.add_page_break()

    # 2. FRONT MATTER
    for section in ["DECLARATION", "ACKNOWLEDGEMENT", "ABSTRACT"]:
        doc.add_heading(section, 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        if section == "ABSTRACT":
            p = doc.add_paragraph(
                "In the current clinical landscape, the lack of health literacy remains a significant barrier to effective diagnosis and self-care. "
                "Patients often receive medical reports filled with high-density jargon that they cannot interpret without professional aid. "
                "The 'Doc 2 Me' system addresses this gap by utilizing state-of-the-art multimodal large language models to translate clinical documents into plain language. "
                "The system integrates Tesseract-based OCR for digitizing hard-copy reports and a Retrieval-Augmented Generation (RAG) pipeline to provide context-aware simplifications. "
                "Key technical contributions include a robust PII (Personally Identifiable Information) scrubbing engine that ensures absolute privacy before data inference. "
                "Expert reviews have validated the output accuracy, marking a major step toward democratic, accessible healthcare information for all demographics."
            )
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            doc.add_paragraph("Content placeholder for " + section.lower() + " section...")
        doc.add_page_break()

    # TOC Placeholders
    for toc in ["TABLE OF CONTENTS", "LIST OF TABLES", "LIST OF FIGURES", "LIST OF ABBREVIATIONS"]:
        doc.add_heading(toc, 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

    # 3. CHAPTER 1
    doc.add_heading('CHAPTER 1', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('INTRODUCTION', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('1.1 Background and Motivation', 2)
    p = doc.add_paragraph(
        "South Asia represents one of the world's most populous regions with a multi-layered healthcare challenge. "
        "While digital health records are proliferating, the 'last mile' of patient understanding is still missing. "
        "Medical jargon, specialized terminology, and complex anatomical references in patient reports often lead to misunderstanding and anxiety."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p = doc.add_paragraph(
        "The consequences of clinical opacity are severe. Patients may ignore critical findings because they do not understand the terminology, "
        "or conversely, experience unnecessary stress over benign findings. The 'Doc 2 Me' project motivation stems from the need to empower "
        "patients with clear, accurate, and private health information summaries."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading('1.2 Problem Statement', 2)
    p = doc.add_paragraph(
        "Existing medical portals provide raw data but no interpretation. General-purpose LLMs lack the domain-specific rigor and privacy "
        "guarantees (PII scrubbing) required for sensitive health tasks. There is an urgent need for a unified platform that can take a physical report, "
        "digitize it, anonymize it, and return a plain-language summary in regional languages."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 4. CHAPTER 2
    doc.add_page_break()
    doc.add_heading('CHAPTER 2', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('LITERATURE SURVEY', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('2.1 Natural Language Processing in Clinical Settings', 2)
    p = doc.add_paragraph(
        "The evolution of NLP from rule-based systems to transformers has revolutionized clinical documentation. "
        "Models like BioBERT (Lee et al., 2020) have set benchmarks for medical named entity recognition (NER). "
        "However, pure classification models are insufficient for abstractive simplification, which requires generative models like T5 and Llama-3."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading('2.2 Summary of Gaps', 2)
    p = doc.add_paragraph("The survey reveals several persistent gaps:")
    for gap in [
        "Lack of Privacy-First Local Inference: Most systems rely on public cloud APIs without local PII redaction.",
        "Monolingual Limitations: Existing tools are predominantly English-only, neglecting regional languages like Nepali.",
        "Absence of Expert Feedback Loop: Most AI models work in a black box without a dedicated medical expert verification mechanism."
    ]:
        doc.add_paragraph(gap, style='List Paragraph')

    # 5. CHAPTER 3
    doc.add_page_break()
    doc.add_heading('CHAPTER 3', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('SYSTEM DESIGN AND ARCHITECTURE', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('3.1 Architecture Overview', 2)
    p = doc.add_paragraph(
        "The Doc 2 Me system is architected as a modular, high-security platform. It follows a multi-tier design Philosophy:"
    )
    for layer in [
        "Interface Layer: Flask/FastAPI based web dashboard for interactive report analysis.",
        "Security Layer: Custom regex-based PII scrubbing engine that anonymizes data locally.",
        "Inference Core: A RAG (Retrieval-Augmented Generation) pipeline connecting LangChain with Llama-3 (8B) and T5 fine-tuned models.",
        "Knowledge Base: FAISS-indexed medical terminology repository for high-fidelity simplification."
    ]:
        doc.add_paragraph(layer, style='List Paragraph')

    # 6. REFERENCES
    doc.add_page_break()
    doc.add_heading('REFERENCES', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("1. Vaswani, A., et al. (2017). 'Attention is All You Need'. NIPS.")
    doc.add_paragraph("2. Devlin, J., et al. (2018). 'BERT: Pre-training of Deep Bidirectional Transformers'.")
    doc.add_paragraph("3. Lewis, M., et al. (2019). 'BART: Denoising Sequence-to-Sequence Pre-training'.")

    # Global Font Fix
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = 'Times New Roman'
    
    output_path = '/Users/kushaldevraj/Downloads/Medical_Jargon_Simplification-main/Doc2Me_Comprehensive_Report.docx'
    doc.save(output_path)
    print(f"Report created at: {output_path}")

if __name__ == '__main__':
    create_comprehensive_report()
