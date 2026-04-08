import sys
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxcompose.composer import Composer

def build_final_document(master_path, sub_path, output_path):
    print("1. Assembling raw merged document...")
    master_doc = Document(master_path)
    
    composer = Composer(master_doc)
    sub_doc = Document(sub_path)
    composer.append(sub_doc)
    
    base_merged_path = output_path.replace('.docx', '_base.docx')
    composer.save(base_merged_path)
    
    print("2. Enhancing document with TOC, LOF, and formatting...")
    doc = Document(base_merged_path)
    
    sectPr = doc.sections[0]._sectPr
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), 'lowerRoman')
    pgNumType.set(qn('w:start'), '1')
    sectPr.append(pgNumType)

    chap1_p = None
    abstract_seen = False
    
    for i, p in enumerate(doc.paragraphs):
        text_lower = p.text.strip().lower()
        if text_lower.startswith('chapter 1') and p.style.name.startswith('Heading'):
            chap1_p = p
            
        if p.style.name == 'Heading 1' and 'abstract' in text_lower:
            if abstract_seen:
                p.clear()
            else:
                abstract_seen = True
                if i > 50:
                    p.insert_paragraph_before().add_run().add_break(WD_BREAK.PAGE)
        elif p.style.name == 'Heading 1':
            p.insert_paragraph_before().add_run().add_break(WD_BREAK.PAGE)

    if chap1_p:
        pb_start = chap1_p.insert_paragraph_before()
        pPr = pb_start._p.get_or_add_pPr()
        sectPr = OxmlElement('w:sectPr')
        
        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(qn('w:fmt'), 'lowerRoman')
        sectPr.append(pgNumType)
        
        type_xml = OxmlElement('w:type')
        type_xml.set(qn('w:val'), 'nextPage')
        sectPr.append(type_xml)
        pPr.append(sectPr)
        
        # Insert TOC Title
        toc_title = chap1_p.insert_paragraph_before('TABLE OF CONTENTS', style='Heading 1')
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Extract Static TOC dynamically
        toc_content = [
            ("Title Page", ""),
            ("Declaration", ""),
            ("Certificate", ""),
            ("Acknowledgement", ""),
            ("Abstract", ""),
            ("Table of Contents", ""),
            ("List of Figures", "")
        ]
        
        # Scan forward for Headings
        for p in doc.paragraphs:
            style = p.style.name
            text = p.text.strip()
            if not text: continue
            
            t_lower = text.lower()
            if style == 'Heading 1' and t_lower not in ['table of contents', 'list of figures', 'abstract', 'co-po-pso mapping', 'sustainable development goals (sdgs)']:
                toc_content.append((text, ""))
            elif style == 'Heading 2':
                toc_content.append((text, ""))
            elif style == 'Heading 3':
                toc_content.append((text, ""))
                
        # Append the ones we inject manually
        toc_content.append(("CO-PO-PSO Mapping", ""))
        toc_content.append(("Sustainable Development Goals (SDGs)", ""))
        
        for item, page in toc_content:
            p = chap1_p.insert_paragraph_before()
            p.style = 'Normal'
            p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            p.add_run(f"{item}\t{page}")

        chap1_p.insert_paragraph_before().add_run().add_break(WD_BREAK.PAGE)
        
        # Insert LOF
        lof_title = chap1_p.insert_paragraph_before('LIST OF FIGURES', style='Heading 1')
        lof_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = chap1_p.insert_paragraph_before()
        p.style = 'Normal'
        p.paragraph_format.tab_stops.add_tab_stop(Inches(1.0), WD_TAB_ALIGNMENT.LEFT)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        p.add_run("Fig No.\tTitle\tPage")
        
        # Extract Static LOF dynamically
        lof_content = []
        fig_idx = 1
        for p in doc.paragraphs:
            if p.style.name == 'Caption' and p.text.strip():
                lof_content.append((f"{fig_idx}.0", p.text.strip(), ""))
                fig_idx += 1
        
        for num, title, page in lof_content:
            p = chap1_p.insert_paragraph_before()
            p.style = 'Normal'
            p.paragraph_format.tab_stops.add_tab_stop(Inches(1.0), WD_TAB_ALIGNMENT.LEFT)
            p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            p.add_run(f"{num}\t{title}\t{page}")
            
        chap1_p.insert_paragraph_before().add_run().add_break(WD_BREAK.PAGE)
        
    doc.add_page_break()
    h = doc.add_paragraph('CO-PO-PSO Mapping', style='Heading 1')
    p1 = doc.add_paragraph('CO–PO–PSO Articulation Table (Project Title: DOC 2 ME: HEALTH REPORT SIMPLIFIER)')
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    headers = ['COs', 'PO1', 'PO2', 'PO3', 'PO4', 'PO5', 'PO6', 'PO7', 'PO8', 'PO9', 'PO10', 'PO11', 'PO12', 'PSO1', 'PSO2', 'PSO3']
    data = [
        ['CO1', '3', '2', '2', '2', '2', '1', '1', '2', '2', '3', '2', '1', '2', '2', '2'],
        ['CO2', '3', '3', '3', '2', '3', '1', '2', '2', '2', '2', '2', '1', '3', '3', '2'],
        ['CO3', '3', '3', '3', '3', '3', '1', '2', '2', '2', '2', '2', '1', '3', '3', '2'],
        ['CO4', '3', '3', '2', '3', '2', '1', '1', '2', '2', '2', '2', '1', '3', '2', '2'],
        ['CO5', '2', '2', '2', '2', '1', '2', '2', '2', '2', '3', '2', '2', '2', '1', '3'],
        ['CO6', '2', '2', '2', '2', '2', '1', '1', '3', '2', '3', '3', '2', '2', '2', '3']
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hc = hdr_cells[i]
        hc.text = title
        hc.paragraphs[0].runs[0].font.bold = True
    for row in data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            
    doc.add_page_break()
    h2 = doc.add_paragraph('Sustainable Development Goals (SDGs)', style='Heading 1')
    sdg_data = [['SDG 3', 'Good Health and Well-being'], ['SDG 4', 'Quality Education'], ['SDG 10', 'Reduced Inequalities']]
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = 'Table Grid'
    t2.rows[0].cells[0].text, t2.rows[0].cells[1].text = 'SDG No.', 'SDG Title'
    for r in sdg_data:
        cells = t2.add_row().cells
        cells[0].text, cells[1].text = r[0], r[1]
        
    p2 = doc.add_paragraph('CO–SDG Mapping with Justification')
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    co_sdg_data = [['CO1', '4', 'Quality Education'], ['CO2', '3', 'Good Health and Well-being'], ['CO3', '3', 'Good Health and Well-being'], ['CO4', '3', 'Good Health and Well-being'], ['CO5', '4', 'Quality Education'], ['CO6', '4', 'Quality Education']]
    t3 = doc.add_table(rows=1, cols=3)
    t3.style = 'Table Grid'
    t3.rows[0].cells[0].text, t3.rows[0].cells[1].text, t3.rows[0].cells[2].text = 'COs', 'SDG No(s)', 'SDG Title(s)'
    for r in co_sdg_data:
        cells = t3.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = r[0], r[1], r[2]
        
    print("3. Enforcing CSE-54 Formatting rules...")
    
    if len(doc.sections) > 1:
        s2 = doc.sections[1]._sectPr
        pgNumType2 = OxmlElement('w:pgNumType')
        pgNumType2.set(qn('w:fmt'), 'decimal')
        pgNumType2.set(qn('w:start'), '1')
        s2.append(pgNumType2)

    for p in doc.paragraphs:
        if not p.text.strip():
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            continue
            
        for run in p.runs:
            run.font.name = 'Times New Roman'
            
        if p.style.name == 'Heading 1':
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(16)
                
        elif p.style.name == 'Heading 2':
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(14)
                
        elif p.style.name.startswith('Heading'):
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(12)
                
        elif p.style.name in ['Normal', 'Body Text'] or p.style.name.startswith('List'):
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0) # Strictly 0 to prevent huge table gaps
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                if not run.font.size:
                    run.font.size = Pt(12)

    doc.save(output_path)
    if os.path.exists(base_merged_path):
        os.remove(base_merged_path)
    print(f"Successfully generated final document at {output_path}")

if __name__ == '__main__':
    base = '/Users/kushaldevraj/Downloads/Medical_Jargon_Simplification-main'
    original_front = os.path.join(base, 'capstonefinal.docx')
    original_chapters = os.path.join(base, 'finalcapstone.docx')
    final_output = os.path.join(base, 'Formatted_Capstone_Report.docx')
    
    build_final_document(original_front, original_chapters, final_output)
