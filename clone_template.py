import sys
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxcompose.composer import Composer

def replace_in_paragraph(p, old_text, new_text):
    if old_text in p.text:
        # A simple replacement that combines all text into the first run to preserve formatting of the start
        if p.runs:
            style = p.style
            for r in p.runs:
                if old_text in r.text:
                    r.text = r.text.replace(old_text, new_text)
                    return True
            # If split across runs, just rewrite the whole paragraph in the first run and clear others
            full_text = p.text
            if old_text in full_text:
                for r in p.runs:
                    r.text = ''
                p.runs[0].text = full_text.replace(old_text, new_text)
                return True
    return False

def clone_and_replace():
    base = '/Users/kushaldevraj/Downloads/Medical_Jargon_Simplification-main'
    template_path = os.path.join(base, 'CSE-54_FRONT PAGE_SAMPLE.docx')
    user_doc_path = os.path.join(base, 'finalcapstone.docx')
    output_path = os.path.join(base, 'Formatted_Capstone_Report.docx')
    
    print("1. Loading Template...")
    doc = Document(template_path)
    
    # 1. STRING REPLACEMENTS IN PARAGRAPHS
    OLD_TITLE = "Hybrid Quantum-Classical Machine Learning for Data Classification"
    NEW_TITLE = "DOC 2 ME: HEALTH REPORT SIMPLIFIER"
    OLD_GUIDE = "Dr Smita Sandeep Darandale"
    NEW_GUIDE = "Sangameshwar"
    OLD_ROLE = "Associate Professor"
    NEW_ROLE = "Assistant Professor"
    
    for p in doc.paragraphs:
        replace_in_paragraph(p, OLD_TITLE, NEW_TITLE)
        replace_in_paragraph(p, "“Hybrid Quantum-Classical Machine Learning for Data Classification”", f"“{NEW_TITLE}”")
        replace_in_paragraph(p, OLD_GUIDE, NEW_GUIDE)
        replace_in_paragraph(p, OLD_ROLE, NEW_ROLE)
        
        # Replace student list string in certificate
        old_students = "Shashank N (BU22CSEN0102266), Shabarish H C (BU22CSEN0102181), Sadananda S V (BU22CSEN0102307), K P Harsha Vardhan (BU22CSEN0100413)"
        new_students = "Kushal Devraj (BU22CSEN0102284)"
        replace_in_paragraph(p, old_students, new_students)
        
    # 2. STRING REPLACEMENTS IN TABLES
    # Table 0: Title page names
    if len(doc.tables) > 0:
        t0 = doc.tables[0]
        for row in t0.rows:
            for cell in row.cells:
                cell.text = ''
        t0.cell(0,0).text = "Kushal Devraj"
        t0.cell(0,1).text = "BU22CSEN0102284"
        
    # Table 1: Declaration names
    if len(doc.tables) > 1:
        t1 = doc.tables[1]
        for row in t1.rows:
            for cell in row.cells:
                cell.text = ''
        t1.cell(0,0).text = "Registration No(s).\nBU22CSEN0102284"
        t1.cell(0,1).text = "Name(s)\nKushal Devraj"
        t1.cell(0,2).text = "Signature(s)\n"
        
    # Table 2: Certificate names
    if len(doc.tables) > 2:
        t2 = doc.tables[2]
        for row in t2.rows:
            for cell in row.cells:
                cell.text = ''
        t2.cell(0,0).text = "Kushal Devraj"
        t2.cell(0,1).text = "BU22CSEN0102284"
        
    # 3. REPLACE ABSTRACT CONTENT
    # Find Abstract Header
    abstract_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == 'ABSTRACT':
            abstract_idx = i
            break
            
    # Delete sample abstract text paragraphs.
    if abstract_idx != -1:
        idx = abstract_idx + 1
        while idx < len(doc.paragraphs):
            p = doc.paragraphs[idx]
            if p.text.strip() == 'TABLE OF CONTENTS' or p.text.strip() == 'LIST OF FIGURES' or p.style.name.startswith('Heading'):
                break
            p._element.getparent().remove(p._element)

        # Insert user's Abstract
        user_doc = Document(user_doc_path)
        user_abs_text = []
        in_abs = False
        for p in user_doc.paragraphs:
            if 'abstract' in p.text.lower() and p.style.name.startswith('Heading'):
                in_abs = True
                continue
            if in_abs:
                if 'chapter 1' in p.text.lower() and p.style.name.startswith('Heading'): break
                if p.text.strip():
                    user_abs_text.append(p.text)
                
        # Insert them backwards so they appear right after ABSTRACT
        for text in reversed(user_abs_text):
            if text.strip():
                if abstract_idx + 1 < len(doc.paragraphs):
                    np = doc.paragraphs[abstract_idx+1].insert_paragraph_before(text)
                else:
                    np = doc.add_paragraph(text)
                np.style = 'Normal'
                np.paragraph_format.line_spacing = 1.5
                np.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 4. DELETE SAMPLE TOC AND LOF TABLES
    # The sample TOC is in Table 3, LOF in Table 4, LOT in Table 5.
    # We will delete them. python-docx doesn't explicitly have table.delete(), so we remove the XML.
    for i in range(len(doc.tables)-1, 2, -1):
        tbl = doc.tables[i]
        tbl._element.getparent().remove(tbl._element)
        
    # Also delete the literal text "TABLE OF CONTENTS", "LIST OF FIGURES", "LIST OF TABLES" from sample
    toc_marker_idx = -1
    for i, p in enumerate(doc.paragraphs):
        low = p.text.strip().lower()
        if low in ['table of contents', 'list of figures', 'list of tables']:
            p.clear()
            
    # Save the prepared front template
    front_prepared_path = os.path.join(base, 'front_prepared.docx')
    doc.save(front_prepared_path)
    
    return front_prepared_path

def build_final_document(master_path, sub_path, output_path):
    print("1. Assembling Document...")
    master_doc = Document(master_path)
    # The template might not end on a clean page break for the appending.
    # Add page break at the end.
    master_doc.add_page_break()
    
    composer = Composer(master_doc)
    sub_doc = Document(sub_path)
    
    # We do NOT append ALL of sub_doc, just from Chapter 1. Wait, docxcompose appends everything.
    # We will just append it and delete the extra Abstract from the composite later.
    composer.append(sub_doc)
    
    base_merged_path = output_path.replace('.docx', '_base.docx')
    composer.save(base_merged_path)
    
    print("2. Enhancing document with TOC, LOF, and formatting...")
    doc = Document(base_merged_path)

    # Find where Chapter 1 starts
    chap1_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if 'chapter 1' in p.text.lower() and p.style.name.startswith('Heading'):
            chap1_idx = i
            break
            
    # Clean up duplicate Abstract from the appended sub_doc
    if chap1_idx != -1:
        abs_start = -1
        for i in range(chap1_idx - 1, -1, -1):
            if 'abstract' in doc.paragraphs[i].text.lower() and doc.paragraphs[i].style.name.startswith('Heading'):
                abs_start = i
                break
        
        if abs_start != -1:
            # Delete everything from chap1_idx-1 down to abs_start safely
            for i in range(chap1_idx - 1, abs_start - 1, -1):
                p = doc.paragraphs[i]
                p._element.getparent().remove(p._element)

    if chap1_idx != -1:
        chap1_p = doc.paragraphs[chap1_idx]
        # Re-apply Section break logic
        pb_start = chap1_p.insert_paragraph_before()
        pPr = pb_start._p.get_or_add_pPr()
        sectPr = OxmlElement('w:sectPr')
        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(qn('w:fmt'), 'lowerRoman')
        sectPr.append(pgNumType)
        type_xml = OxmlElement('w:type')
        type_xml.set(qn('w:val'), 'nextPage')
        sectPr.append(type_xml)
        pPr.append(sectPr)
        
        # Build Static TOC
        toc_title = chap1_p.insert_paragraph_before('TABLE OF CONTENTS', style='Heading 1')
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        toc_content = [
            ("Title Page", ""),
            ("Declaration", ""),
            ("Certificate", ""),
            ("Acknowledgement", ""),
            ("Abstract", ""),
            ("Table of Contents", ""),
            ("List of Figures", "")
        ]
        
        for p in doc.paragraphs:
            style = p.style.name
            text = p.text.strip()
            if not text: continue
            
            t_lower = text.lower()
            if style == 'Heading 1' and t_lower not in ['table of contents', 'list of figures', 'abstract', 'co-po-pso mapping', 'sustainable development goals (sdgs)']:
                toc_content.append((text, ""))
            elif style == 'Heading 2':
                toc_content.append((text, ""))
            elif style == 'Heading 3':
                toc_content.append((text, ""))
                
        toc_content.append(("CO-PO-PSO Mapping", ""))
        toc_content.append(("Sustainable Development Goals (SDGs)", ""))
        
        for item, page in toc_content:
            p = chap1_p.insert_paragraph_before()
            p.style = 'Normal'
            p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            p.add_run(f"{item}\t{page}")
        
        # Build Static LOF
        lof_title = chap1_p.insert_paragraph_before('LIST OF FIGURES', style='Heading 1')
        lof_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = chap1_p.insert_paragraph_before()
        p.style = 'Normal'
        p.paragraph_format.tab_stops.add_tab_stop(Inches(1.0), WD_TAB_ALIGNMENT.LEFT)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        p.add_run("Fig No.\tTitle\tPage")
        
        lof_content = []
        fig_idx = 1
        for p in doc.paragraphs:
            if p.style.name == 'Caption' and p.text.strip():
                lof_content.append((f"{fig_idx}.0", p.text.strip(), ""))
                fig_idx += 1
        
        for num, title, page in lof_content:
            p = chap1_p.insert_paragraph_before()
            p.style = 'Normal'
            p.paragraph_format.tab_stops.add_tab_stop(Inches(1.0), WD_TAB_ALIGNMENT.LEFT)
            p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            p.add_run(f"{num}\t{title}\t{page}")

    # Append CO-PO Mappings
    h = doc.add_paragraph('CO-PO-PSO Mapping', style='Heading 1')
    p1 = doc.add_paragraph('CO–PO–PSO Articulation Table (Project Title: DOC 2 ME: HEALTH REPORT SIMPLIFIER)')
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    headers = ['COs', 'PO1', 'PO2', 'PO3', 'PO4', 'PO5', 'PO6', 'PO7', 'PO8', 'PO9', 'PO10', 'PO11', 'PO12', 'PSO1', 'PSO2', 'PSO3']
    data = [
        ['CO1', '3', '2', '2', '2', '2', '1', '1', '2', '2', '3', '2', '1', '2', '2', '2'],
        ['CO2', '3', '3', '3', '2', '3', '1', '2', '2', '2', '2', '2', '1', '3', '3', '2'],
        ['CO3', '3', '3', '3', '3', '3', '1', '2', '2', '2', '2', '2', '1', '3', '3', '2'],
        ['CO4', '3', '3', '2', '3', '2', '1', '1', '2', '2', '2', '2', '1', '3', '2', '2'],
        ['CO5', '2', '2', '2', '2', '1', '2', '2', '2', '2', '3', '2', '2', '2', '1', '3'],
        ['CO6', '2', '2', '2', '2', '2', '1', '1', '3', '2', '3', '3', '2', '2', '2', '3']
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = 'Table Grid'
    except KeyError:
        pass # The template does not possess the default Table Grid style.
        
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hc = hdr_cells[i]
        hc.text = title
        hc.paragraphs[0].runs[0].font.bold = True
    for row in data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            
    h2 = doc.add_paragraph('Sustainable Development Goals (SDGs)', style='Heading 1')
    sdg_data = [['SDG 3', 'Good Health and Well-being'], ['SDG 4', 'Quality Education'], ['SDG 10', 'Reduced Inequalities']]
    t2 = doc.add_table(rows=1, cols=2)
    try:
        t2.style = 'Table Grid'
    except KeyError:
        pass
    t2.rows[0].cells[0].text, t2.rows[0].cells[1].text = 'SDG No.', 'SDG Title'
    for r in sdg_data:
        cells = t2.add_row().cells
        cells[0].text, cells[1].text = r[0], r[1]
        
    p2 = doc.add_paragraph('CO–SDG Mapping with Justification')
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    co_sdg_data = [['CO1', '4', 'Quality Education'], ['CO2', '3', 'Good Health and Well-being'], ['CO3', '3', 'Good Health and Well-being'], ['CO4', '3', 'Good Health and Well-being'], ['CO5', '4', 'Quality Education'], ['CO6', '4', 'Quality Education']]
    t3 = doc.add_table(rows=1, cols=3)
    try:
        t3.style = 'Table Grid'
    except KeyError:
        pass
    t3.rows[0].cells[0].text, t3.rows[0].cells[1].text, t3.rows[0].cells[2].text = 'COs', 'SDG No(s)', 'SDG Title(s)'
    for r in co_sdg_data:
        cells = t3.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = r[0], r[1], r[2]

    # Enforce global rules on ALL paragraphs just to be perfectly compliant
    print("3. Enforcing Global Formatting Rules...")
    if len(doc.sections) > 1:
        s2 = doc.sections[1]._sectPr
        pgNumType2 = OxmlElement('w:pgNumType')
        pgNumType2.set(qn('w:fmt'), 'decimal')
        pgNumType2.set(qn('w:start'), '1')
        s2.append(pgNumType2)

    for p in doc.paragraphs:
        if not p.text.strip():
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            continue
            
        for run in p.runs:
            run.font.name = 'Times New Roman'
            
        if p.style.name == 'Heading 1':
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.page_break_before = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(16)
        
        # Don't ruin front page formatting
        # Only touch Normal Text
        elif p.style.name in ['Normal', 'Body Text'] and not p.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                if not run.font.size: run.font.size = Pt(12)

    doc.save(output_path)
    if os.path.exists(base_merged_path):
        os.remove(base_merged_path)
    print(f"Successfully generated final document at {output_path}")

if __name__ == '__main__':
    base = '/Users/kushaldevraj/Downloads/Medical_Jargon_Simplification-main'
    user_chapters = os.path.join(base, 'finalcapstone.docx')
    final_output = os.path.join(base, 'Formatted_Capstone_Report.docx')
    
    front_prepared = clone_and_replace()
    build_final_document(front_prepared, user_chapters, final_output)
