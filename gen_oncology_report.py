from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_oncology_report():
    base_path = "/Users/kushaldevraj/Downloads/Medical_Jargon_Simplification-main/"
    filename_base = "Sample_Oncology_Pathology_Report"
    
    # 1. DATA
    header_title = "HISTOPATHOLOGY EXAMINATION REPORT"
    category = "Oncology / Surgical Pathology"
    orig_text = "Specimen obtained via ultrasound-guided core needle biopsy of retroareolar mass."
    results_data = [
        ("Histological Type", "Invasive Ductal Carcinoma (IDC)", "", "Benign"),
        ("Nottingham Grade", "Grade 2 (Intermediate)", "Score: 6/9", "Low Grade"),
        ("ER / PR Status", "Estrogen Positive (80%)", "", "N/A"),
        ("HER2/neu", "3+ (Positive)", "", "Negative"),
        ("TNM Staging", "pT2 N1 M0", "", "Early Stage")
    ]
    impression_text = (
        "Malignant findings consistent with Hormone Receptor-positive, HER2-enriched Invasive Ductal Carcinoma. "
        "Evidence of lymphovascular invasion noted in sentinel node biopsied. Stage IIA disease. "
        "Multidisciplinary tumor board review and oncological consultation for systemic therapy recommended."
    )

    # 2. DOCX
    doc = Document()
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(f"REGIONAL ONCOLOGY CENTER - {category.upper()}\n")
    run.font.size = Pt(16)
    run.font.bold = True
    doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(header_title, 2)
    
    results_table = doc.add_table(rows=1, cols=4)
    results_table.style = 'Table Grid'
    hdr = results_table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "DIAGNOSTIC PARAMETER", "FINDING", "DETAIL", "REFERENCE"
    for p, r, u, ref in results_data:
        row = results_table.add_row().cells
        row[0].text, row[1].text, row[2].text, row[3].text = p, str(r), u, ref
    
    doc.add_paragraph("\n")
    imp = doc.add_paragraph()
    imp.add_run("CLINICAL IMPRESSION: ").font.bold = True
    imp.add_run(impression_text)
    doc.save(f"{base_path}{filename_base}.docx")

    # 3. JPEG
    width, height = 2480, 3508
    image = Image.new('RGB', (width, height), color='white')
    draw = ImageDraw.Draw(image)
    try:
        font_bold = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 80)
        font_reg = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 60)
        font_header = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 100)
    except: font_bold = font_reg = font_header = ImageFont.load_default()
    
    margin, y = 150, 200
    draw.text((width/2, y), "REGIONAL ONCOLOGY CENTER", fill='black', font=font_header, anchor="mm")
    y += 150
    draw.line((margin, y, width-margin, y), fill='black', width=5)
    y += 150
    draw.text((margin, y), header_title, fill='black', font=font_header)
    y += 200
    for p, r, u, ref in results_data:
        draw.text((margin, y), f"{p}: {r}", fill='black', font=font_reg)
        y += 120
    y += 150
    draw.text((margin, y), "CLINICAL IMPRESSION:", fill='black', font=font_bold)
    y += 120
    # Wrapping
    words = impression_text.split()
    line = ""
    for word in words:
        if len(line + word) < 55: line += word + " "
        else:
            draw.text((margin + 50, y), line, fill='black', font=font_reg)
            y += 80
            line = word + " "
    draw.text((margin + 50, y), line, fill='black', font=font_reg)
    image.save(f"{base_path}{filename_base}.jpg")
    print(f"Created Oncology assets at {base_path}{filename_base}")

if __name__ == '__main__':
    create_oncology_report()
