
import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_document_font(doc, font_name='Times New Roman'):
    """Set the default font for the document styles"""
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(12)
    
    # Also set for headings
    for i in range(1, 4):
        style = doc.styles[f'Heading {i}']
        style.font.name = font_name
        style.font.color.rgb = RGBColor(0, 0, 0) # Black headings
        
    # Ensure all runs use the font
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:ascii'), font_name)
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)

def parse_markdown(md_text, doc):
    lines = md_text.split('\n')
    in_code_block = False
    code_content = []
    
    # Simple table parser state
    in_table = False
    table_headers = []
    table_rows = []

    for line in lines:
        line = line.strip()
        
        # Code block handling
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_content))
                run.font.name = 'Courier New' # Keep code monospaced
                run.font.size = Pt(10)
                p.style = 'No Spacing' 
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            continue
            
        if in_code_block:
            code_content.append(line)
            continue

        # Image Handling: ![Caption](path)
        img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
        if img_match:
            caption = img_match.group(1)
            img_path = img_match.group(2)
            
            # Resolve path relative to current directory if needed
            if not os.path.isabs(img_path):
                img_path = os.path.abspath(img_path)
                
            if os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(6.0))
                    last_paragraph = doc.paragraphs[-1] 
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    if caption:
                        caption_p = doc.add_paragraph(caption)
                        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_p.style = 'Caption'
                        caption_run = caption_p.runs[0]
                        caption_run.font.name = 'Times New Roman'
                        caption_run.italic = True
                except Exception as e:
                    print(f"Error adding image {img_path}: {e}")
                    doc.add_paragraph(f"[Image: {caption} - Failed to load]")
            else:
                doc.add_paragraph(f"[Image: {caption} - File not found: {img_path}]")
            continue

        # Table handling (basic)
        if line.startswith('|') and line.endswith('|'):
            if '---' in line: # Separator line
                continue
            
            cells = [c.strip() for c in line.split('|') if c]
            if not in_table:
                in_table = True
                table_headers = cells
            else:
                table_rows.append(cells)
            continue
        elif in_table:
            # End of table, render it now
            table = doc.add_table(rows=1, cols=len(table_headers))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(table_headers):
                if i < len(hdr_cells):
                    hdr_cells[i].text = h
                    # varied styling for header could go here
            
            for row_data in table_rows:
                row = table.add_row().cells
                for i, cell_data in enumerate(row_data):
                    if i < len(row):
                        row[i].text = cell_data
            
            doc.add_paragraph() # Spacer
            in_table = False
            table_headers = []
            table_rows = []
        
        if not line:
            continue

        # Headers
        if line.startswith('# '):
            h = doc.add_heading(line[2:], level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        
        # Lists
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, line[2:])
        elif re.match(r'^\d+\.\s', line):
            # Numbered list
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
            
        # Blockquotes/Notes
        elif line.startswith('> '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.italic = True
            run.font.name = 'Times New Roman'
            p.paragraph_format.left_indent = Inches(0.5)

        # Standard Paragraph
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, line)

    # Catch trailing table if file ends with one
    if in_table:
        table = doc.add_table(rows=1, cols=len(table_headers))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(table_headers):
            if i < len(hdr_cells):
                 hdr_cells[i].text = h
        for row_data in table_rows:
            row = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                if i < len(row):
                    row[i].text = cell_data


def add_formatted_text(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Times New Roman'
        else:
            # Handle italics inside non-bold parts
            italic_parts = re.split(r'(\*.*?\*)', part)
            for ipart in italic_parts:
                if ipart.startswith('*') and ipart.endswith('*') and len(ipart) > 2:
                    run = paragraph.add_run(ipart[1:-1])
                    run.italic = True
                    run.font.name = 'Times New Roman'
                else:
                    run = paragraph.add_run(ipart)
                    run.font.name = 'Times New Roman'

if __name__ == "__main__":
    try:
        with open('PROJECT_REPORT_FINAL.md', 'r', encoding='utf-8') as f:
            content = f.read()
            
        document = Document()
        set_document_font(document, 'Times New Roman')
        
        document.add_heading('MediClare - Medical Jargon Simplification Report', 0)
        
        parse_markdown(content, document)
        
        output_file = 'MediClare_Project_Report_v2.docx'
        document.save(output_file)
        print(f"Successfully created {output_file}")
    except Exception as e:
        print(f"Error: {e}")
