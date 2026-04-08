import sys
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxcompose.composer import Composer

def set_font_times_new_roman_and_justify(doc):
    """Enforce Times New Roman font and justified alignment for all runs."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    abstract_seen = False
    
    for p in doc.paragraphs:
        # Enforce Times New Roman on every run
        for run in p.runs:
            run.font.name = 'Times New Roman'
        
        # Enforce Justified alignment for Normal text
        if p.style.name == 'Normal' or p.style.name == 'Body Text':
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
        # Ensure Heading 1 starts on a new page
        text_lower = p.text.strip().lower()
        if p.style.name == 'Heading 1':
            # Remove duplicate abstracts if they somehow got in
            if 'abstract' in text_lower:
                if abstract_seen:
                    p.clear()  # Delete duplicate abstract heading
                    continue
                abstract_seen = True
                
            for run in p.runs:
                run.font.bold = True

def insert_field(p, code):
    """Insert Word field code into a paragraph."""
    run = p.add_run()
    fld1 = OxmlElement('w:fldChar')
    fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = code
    fld2 = OxmlElement('w:fldChar')
    fld2.set(qn('w:fldCharType'), 'separate')
    fld3 = OxmlElement('w:fldChar')
    fld3.set(qn('w:fldCharType'), 'end')
    for f in (fld1, instr, fld2, fld3):
        run._r.append(f)

def prepare_finalcapstone(input_path, output_path):
    doc = Document(input_path)
    
    # Force page breaks before all chapters and abstract
    for p in doc.paragraphs:
        if p.style.name == 'Heading 1' and p.text.strip().lower().startswith('chapter '):
            p.insert_paragraph_before().add_run().add_break(WD_BREAK.PAGE)
    
    # Find where Chapter 1 starts to insert TOC and LOF before it
    chap1_p = None
    for p in doc.paragraphs:
        if p.text.strip().lower().startswith('chapter 1'):
            chap1_p = p
            break
            
    if chap1_p:
        # Insert Page Break before TOC
        pb_start = chap1_p.insert_paragraph_before()
        pb_start.add_run().add_break(WD_BREAK.PAGE)
        
        # Insert TOC
        toc_title = chap1_p.insert_paragraph_before('TABLE OF CONTENTS', style='Heading 1')
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_instruct = chap1_p.insert_paragraph_before('(Microsoft Word: Right-click below and select "Update Field" to generate Table of Contents)')
        toc_instruct.style.font.italic = True
        toc_field = chap1_p.insert_paragraph_before()
        insert_field(toc_field, 'TOC \\o "1-3" \\h \\z \\u')
        
        # Page Break
        chap1_p.insert_paragraph_before().add_run().add_break(WD_BREAK.PAGE)
        
        # Insert LOF
        lof_title = chap1_p.insert_paragraph_before('LIST OF FIGURES', style='Heading 1')
        lof_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lof_instruct = chap1_p.insert_paragraph_before('(Microsoft Word: Right-click below and select "Update Field" to generate List of Figures)')
        lof_instruct.style.font.italic = True
        lof_field = chap1_p.insert_paragraph_before()
        insert_field(lof_field, 'TOC \\h \\z \\c "Figure"')
        
        # Page Break
        chap1_p.insert_paragraph_before().add_run().add_break(WD_BREAK.PAGE)

    # Add CO-PO-PSO and SDGs at the end of the document
    doc.add_page_break()
    h = doc.add_paragraph('CO-PO-PSO Mapping', style='Heading 1')
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('CO–PO–PSO Articulation Table (Project Title: DOC 2 ME: HEALTH REPORT SIMPLIFIER)')
    headers = ['COs', 'PO1', 'PO2', 'PO3', 'PO4', 'PO5', 'PO6', 'PO7', 'PO8', 'PO9', 'PO10', 'PO11', 'PO12', 'PSO1', 'PSO2', 'PSO3']
    data = [
        ['CO1', '3', '2', '2', '2', '2', '1', '1', '2', '2', '3', '2', '1', '2', '2', '2'],
        ['CO2', '3', '3', '3', '2', '3', '1', '2', '2', '2', '2', '2', '1', '3', '3', '2'],
        ['CO3', '3', '3', '3', '3', '3', '1', '2', '2', '2', '2', '2', '1', '3', '3', '2'],
        ['CO4', '3', '3', '2', '3', '2', '1', '1', '2', '2', '2', '2', '1', '3', '2', '2'],
        ['CO5', '2', '2', '2', '2', '1', '2', '2', '2', '2', '3', '2', '2', '2', '1', '3'],
        ['CO6', '2', '2', '2', '2', '2', '1', '1', '3', '2', '3', '3', '2', '2', '2', '3'],
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hc = hdr_cells[i]
        hc.text = title
        hc.paragraphs[0].runs[0].font.bold = True
    for row in data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            
    doc.add_page_break()
    h2 = doc.add_paragraph('Sustainable Development Goals (SDGs)', style='Heading 1')
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sdg_data = [['SDG 3', 'Good Health and Well-being'], ['SDG 4', 'Quality Education'], ['SDG 10', 'Reduced Inequalities']]
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = 'Table Grid'
    t2.rows[0].cells[0].text, t2.rows[0].cells[1].text = 'SDG No.', 'SDG Title'
    for r in sdg_data:
        cells = t2.add_row().cells
        cells[0].text, cells[1].text = r[0], r[1]
        
    doc.add_paragraph()
    co_sdg_data = [['CO1', '4', 'Quality Education'], ['CO2', '3', 'Good Health and Well-being'], ['CO3', '3', 'Good Health and Well-being'], ['CO4', '3', 'Good Health and Well-being'], ['CO5', '4', 'Quality Education'], ['CO6', '4', 'Quality Education']]
    t3 = doc.add_table(rows=1, cols=3)
    t3.style = 'Table Grid'
    t3.rows[0].cells[0].text, t3.rows[0].cells[1].text, t3.rows[0].cells[2].text = 'COs', 'SDG No(s)', 'SDG Title(s)'
    for r in co_sdg_data:
        cells = t3.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = r[0], r[1], r[2]
        
    set_font_times_new_roman_and_justify(doc)
    doc.save(output_path)
    print(f"Prepared finalcapstone at {output_path}")

def merge_documents(master_path, sub_path, output_path):
    print("Merging documents...")
    master_doc = Document(master_path)
    set_font_times_new_roman_and_justify(master_doc)
    # The last paragraph of master is usually Acknowledgements. Let's add a page break so Abstract starts fresh
    master_doc.add_page_break()
    composer = Composer(master_doc)
    sub_doc = Document(sub_path)
    composer.append(sub_doc)
    set_font_times_new_roman_and_justify(composer.doc)
    composer.save(output_path)
    print(f"Successfully merged documents into {output_path}")

if __name__ == '__main__':
    base = '/Users/kushaldevraj/Downloads/Medical_Jargon_Simplification-main'
    original_front = os.path.join(base, 'capstonefinal.docx')
    original_chapters = os.path.join(base, 'finalcapstone.docx')
    prepared_chapters = os.path.join(base, 'finalcapstone_prepared.docx')
    final_output = os.path.join(base, 'Formatted_Capstone_Report.docx')
    
    prepare_finalcapstone(original_chapters, prepared_chapters)
    merge_documents(original_front, prepared_chapters, final_output)
    print("Done!")
