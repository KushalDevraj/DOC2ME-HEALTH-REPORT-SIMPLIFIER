from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_sample_docx():
    doc = Document()
    
    # 1. Header
    header = doc.add_paragraph() # Use add_paragraph instead of accessing index 0
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("CITY DIAGNOSTIC CENTER & PATHOLOGY LAB\n")
    run.font.size = Pt(16)
    run.font.bold = True
    
    sub = doc.add_paragraph("123 Medical Drive, Health City | Ph: +91-9988776655")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2. Patient Info
    table = doc.add_table(rows=3, cols=2)
    table.columns[0].width = Inches(1.5)
    
    table.cell(0, 0).text = "PATIENT NAME:"
    table.cell(0, 1).text = "MR. KUSHAL DEVRAJ"
    table.cell(1, 0).text = "AGE/SEX:"
    table.cell(1, 1).text = "24 Years / Male"
    table.cell(2, 0).text = "DATE:"
    table.cell(2, 1).text = "2026-03-30"

    doc.add_paragraph("\n")

    # 3. Test Body
    doc.add_heading("TEST: COMPLETE BLOOD COUNT (CBC)", 2)
    
    results = doc.add_table(rows=1, cols=4)
    results.style = 'Table Grid'
    hdr = results.rows[0].cells
    hdr[0].text = "PARAMETER"
    hdr[1].text = "RESULT"
    hdr[2].text = "UNITS"
    hdr[3].text = "REF. RANGE"
    
    data = [
        ("WBC Count", "18.5 *", "10^3/uL", "4.0 - 11.0"),
        ("Neutrophils", "85 *", "%", "40 - 75"),
        ("Hemoglobin", "10.2 *", "g/dL", "13.0 - 17.0"),
        ("Platelet Count", "82 *", "10^3/uL", "150 - 450"),
        ("MCV", "78", "fL", "80 - 100")
    ]
    
    for p, r, u, ref in data:
        row = results.add_row().cells
        row[0].text = p
        row[1].text = r
        row[2].text = u
        row[3].text = ref

    doc.add_paragraph("\n")

    # 4. Impression
    imp_hdr = doc.add_paragraph()
    run = imp_hdr.add_run("IMPRESSION & CLINICAL CORRELATION:")
    run.font.bold = True
    
    doc.add_paragraph(
        "Marked Leukocytosis with Neutrophilia (Left Shift) noted. "
        "Thrombocytopenia present. Microcytic Anemia observed. "
        "Findings are suggestive of an acute systemic infection. Clinical correlation with CRP and fever profile is recommended."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 5. Signature
    sig = doc.add_paragraph("\n\nDr. Pathologist, MBBS, MD")
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    filename = "/Users/kushaldevraj/Downloads/Medical_Jargon_Simplification-main/Sample_Hematology_Report.docx"
    doc.save(filename)
    print(f"Created {filename}")

if __name__ == '__main__':
    create_sample_docx()
