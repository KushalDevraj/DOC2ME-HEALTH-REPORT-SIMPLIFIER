#!/usr/bin/env python3
"""
Fix the corrupted capstonefinal.docx by extracting file contents from local 
file headers (ignoring the corrupted central directory), then rebuild a valid ZIP.
"""
import struct
import zlib
import zipfile
import io
import os
import sys

def parse_local_file_entry(data, offset, next_offset=None):
    """Parse a single local file header and return (name, uncompressed_data) or None."""
    # Local file header: PK\x03\x04
    if data[offset:offset+4] != b'PK\x03\x04':
        return None

    header = data[offset+4:offset+30]
    if len(header) < 26:
        return None

    (version, flags, method, mod_time, mod_date, crc32_val,
     comp_size, uncomp_size, name_len, extra_len) = struct.unpack('<HHHHHIIIHH', header)

    name_start = offset + 30
    name_bytes = data[name_start:name_start + name_len]
    try:
        file_name = name_bytes.decode('utf-8')
    except:
        file_name = name_bytes.decode('latin-1')

    data_start = name_start + name_len + extra_len

    # Determine compressed size
    has_descriptor = (flags & 0x08) != 0

    if comp_size > 0 and not has_descriptor:
        file_data = data[data_start:data_start + comp_size]
    elif next_offset is not None:
        # Use the gap between this entry and the next header
        # But we need to account for possible data descriptor
        raw_region = data[data_start:next_offset]
        
        # Try to find data descriptor in the region
        # Data descriptor can be: PK\x07\x08 + crc32 + comp_size + uncomp_size (16 bytes)
        # or just crc32 + comp_size + uncomp_size (12 bytes)
        dd_sig = raw_region.rfind(b'PK\x07\x08')
        if dd_sig != -1 and dd_sig > 0:
            file_data = raw_region[:dd_sig]
            desc = raw_region[dd_sig+4:dd_sig+16]
            if len(desc) >= 12:
                crc32_val, comp_size, uncomp_size = struct.unpack('<III', desc)
                file_data = raw_region[:comp_size]
        else:
            file_data = raw_region
    else:
        # Last file before central directory
        cd_start = data.find(b'PK\x01\x02', data_start)
        if cd_start == -1:
            # No central directory? Try EOCD
            cd_start = data.find(b'PK\x05\x06', data_start)
        if cd_start == -1:
            cd_start = len(data)
        file_data = data[data_start:cd_start]

    # Decompress
    if method == 8:  # DEFLATED
        try:
            decompressed = zlib.decompress(file_data, -15)
        except zlib.error:
            try:
                decompressed = zlib.decompress(file_data)
            except zlib.error:
                # Try with raw deflate for various window sizes
                for wbits in [-15, -14, -13, 15, 31, 47]:
                    try:
                        decompressed = zlib.decompress(file_data, wbits)
                        break
                    except:
                        continue
                else:
                    print(f"  WARNING: Cannot decompress '{file_name}', skipping...")
                    return None
    elif method == 0:  # STORED
        decompressed = file_data
    else:
        print(f"  WARNING: Unknown compression method {method} for '{file_name}', skipping...")
        return None

    return (file_name, decompressed)


def rebuild_docx(input_path, output_path):
    """Rebuild the DOCX by extracting from local file headers."""
    with open(input_path, 'rb') as f:
        data = f.read()

    print(f"Input file: {len(data)} bytes")

    # Find all local file headers
    positions = []
    pos = 0
    while True:
        pos = data.find(b'PK\x03\x04', pos)
        if pos == -1:
            break
        positions.append(pos)
        pos += 4

    print(f"Found {len(positions)} local file headers")

    # Parse entries
    entries = {}
    for i, hdr_pos in enumerate(positions):
        next_pos = positions[i+1] if i+1 < len(positions) else None
        result = parse_local_file_entry(data, hdr_pos, next_pos)
        if result:
            name, content = result
            # Skip duplicates (keep first occurrence) unless it's empty
            if name not in entries or len(entries[name]) == 0:
                entries[name] = content
            print(f"  Extracted: {name} ({len(content)} bytes)")
        else:
            print(f"  Skipped entry at offset {hdr_pos}")

    # Verify we have the essential docx files
    essential = ['[Content_Types].xml', 'word/document.xml', '_rels/.rels']
    for e in essential:
        if e not in entries:
            print(f"  ERROR: Missing essential file: {e}")
            # Try case-insensitive or partial match
            for k in entries:
                if e.lower() in k.lower():
                    print(f"    Found possible match: {k}")

    # Write the rebuilt ZIP
    output_buf = io.BytesIO()
    with zipfile.ZipFile(output_buf, 'w', zipfile.ZIP_DEFLATED) as zout:
        for name, content in entries.items():
            zout.writestr(name, content)

    with open(output_path, 'wb') as f:
        f.write(output_buf.getvalue())

    print(f"\nRebuilt file: {output_path} ({os.path.getsize(output_path)} bytes)")
    return output_path


def merge_documents(master_path, sub_path, output_path):
    """Merge two docx files preserving formatting."""
    from docx import Document
    from docxcompose.composer import Composer

    master_doc = Document(master_path)
    composer = Composer(master_doc)
    sub_doc = Document(sub_path)
    composer.append(sub_doc)
    composer.save(output_path)
    print(f"\nMerged file saved to: {output_path}")


if __name__ == '__main__':
    base = '/Users/kushaldevraj/Downloads/Medical_Jargon_Simplification-main'
    corrupted = os.path.join(base, 'capstonefinal.docx')
    repaired = os.path.join(base, 'capstonefinal_repaired.docx')
    second = os.path.join(base, 'finalcapstone.docx')
    merged = os.path.join(base, 'merged_capstone.docx')

    print("=" * 60)
    print("STEP 1: Repairing capstonefinal.docx")
    print("=" * 60)
    rebuild_docx(corrupted, repaired)

    print("\n" + "=" * 60)
    print("STEP 2: Verifying repaired file")
    print("=" * 60)
    try:
        from docx import Document
        doc = Document(repaired)
        print(f"SUCCESS! Paragraphs: {len(doc.paragraphs)}, "
              f"Tables: {len(doc.tables)}, "
              f"Sections: {len(doc.sections)}")
    except Exception as e:
        print(f"Verification failed: {e}")
        sys.exit(1)

    print("\n" + "=" * 60)
    print("STEP 3: Merging with finalcapstone.docx")
    print("=" * 60)
    merge_documents(repaired, second, merged)
    print(f"\n✅ Done! Your merged file is at:\n   {merged}")
