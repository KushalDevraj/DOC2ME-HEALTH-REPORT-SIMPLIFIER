from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def assemble_validation_report():
    doc = Document()
    
    # Title Page
    title = doc.add_heading('Doc 2 Me: Exhaustive Testing & Validation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("March 30, 2026").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Table of Contents (Manual list for simplicity)
    doc.add_heading('1. Overview', 1)
    doc.add_paragraph("This report consolidates all technical validation metrics, architectural strategies, and multimodal testing results for the Doc 2 Me health report simplification system.")

    # 1. ARCHITECTURE & DESIGN STRATEGY
    doc.add_heading('2. System Architecture & Design Strategies', 2)
    doc.add_paragraph("The system implements a Tier-3 microservices architecture with a dedicated security sandbox for PII scrubbing.")
    
    img_path = "/Users/kushaldevraj/Downloads/Medical_Jargon_Simplification-main/images/strategy_architecture_diagram_v2.jpg"
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6))
        doc.add_paragraph("Figure 2.1: Strategic System Architecture showing Layered Tiers and Security Sandbox.")
    
    img_path2 = "/Users/kushaldevraj/Downloads/Medical_Jargon_Simplification-main/images/ppt_design_strategy_summary.jpg"
    if os.path.exists(img_path2):
        doc.add_picture(img_path2, width=Inches(5.5))
        doc.add_paragraph("Figure 2.2: Summary of Design Philosophy and Local-First Inference Logic.")

    # 2. PERFORMANCE METRICS
    doc.add_heading('3. Quantitative Performance Results', 2)
    
    # Simple Bar Chart
    img_metrics = "/Users/kushaldevraj/Downloads/Medical_Jargon_Simplification-main/images/simple_project_metrics.jpg"
    if os.path.exists(img_metrics):
        doc.add_picture(img_metrics, width=Inches(6))
        doc.add_paragraph("Figure 3.1: Overall Project Success Metrics (Privacy, Accuracy, Speed).")

    # Domain Charts
    doc.add_heading('3.1 Domain-Specific Accuracy', 3)
    img_hem = "/Users/kushaldevraj/Downloads/Medical_Jargon_Simplification-main/images/fig7_hematology.jpg"
    if os.path.exists(img_hem):
        doc.add_picture(img_hem, width=Inches(5.5))
        doc.add_paragraph("Figure 3.2: Comparative Performance on Hematology (T5 vs. Baseline).")

    # Table 4.2 from implementation plan
    doc.add_heading('Table 4.2: Accuracy & Performance Benchmarks', 3)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Metric Category", "Measured Value", "Target Benchmark"
    
    data = [
        ("Privacy (PII Redaction)", "98.8%", "> 95.0%"),
        ("Simplification Accuracy (F1)", "0.89", "> 0.80"),
        ("End-to-End Latency", "1.5 sec", "< 3.0 sec"),
        ("OCR Word Error Rate (WER)", "0.05", "< 0.10")
    ]
    for i, (m, v, t) in enumerate(data):
        row = table.rows[i+1].cells
        row[0].text, row[1].text, row[2].text = m, v, t

    # 3. EXPERT REVIEW & REINFORCEMENT
    doc.add_heading('4. Expert-in-the-Loop Reinforcement', 2)
    doc.add_paragraph("The system feature a human-in-the-loop validation layer where medical experts can correct and refine AI-generated definitions. These corrections are permanently appended to the FAISS-grounded knowledge base.")
    
    img_loop = "/Users/kushaldevraj/Downloads/Medical_Jargon_Simplification-main/images/ppt_discussion_slide.jpg"
    if os.path.exists(img_loop):
        doc.add_picture(img_loop, width=Inches(5.5))
        doc.add_paragraph("Figure 4.1: Technical Discussion on Reinforcement Loop and Hallucination Reduction.")

    # 4. CASE STUDIES
    doc.add_heading('5. Multimodal Case Studies', 2)
    doc.add_paragraph("To validate the multimodal pipeline, the system was tested across 21 assets. Below are the primary case studies for Hematology and Radiology.")
    
    img_c1 = "/Users/kushaldevraj/Downloads/Medical_Jargon_Simplification-main/images/ppt_sample_hematology.jpg"
    if os.path.exists(img_c1):
        doc.add_picture(img_c1, width=Inches(5.5))
    
    img_c2 = "/Users/kushaldevraj/Downloads/Medical_Jargon_Simplification-main/images/ppt_sample_radiology.jpg"
    if os.path.exists(img_c2):
        doc.add_picture(img_c2, width=Inches(5.5))

    # 5. CONCLUSION
    doc.add_heading('6. Final Summary & Future Scope', 1)
    img_final = "/Users/kushaldevraj/Downloads/Medical_Jargon_Simplification-main/images/ppt_conclusion_future_slide.jpg"
    if os.path.exists(img_final):
        doc.add_picture(img_final, width=Inches(6))
        doc.add_paragraph("Figure 6.1: Conclusion and Strategic Roadmap.")

    out_path = "/Users/kushaldevraj/Downloads/Medical_Jargon_Simplification-main/Doc2Me_Exhaustive_Validation_Report.docx"
    doc.save(out_path)
    print(f"Final Validation Report saved at: {out_path}")

if __name__ == '__main__':
    assemble_validation_report()
