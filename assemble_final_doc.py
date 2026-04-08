import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_final_document():
    doc = Document()
    img_dir = '/Users/kushaldevraj/Downloads/Medical_Jargon_Simplification-main/images'
    
    # 1. TITLE PAGE
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("\n\n\nDOC 2 ME: HEALTH REPORT SIMPLIFIER USING MULTIMODAL LARGE LANGUAGE MODELS\n\n\n")
    run.font.size = Pt(22)
    run.font.bold = True
    
    run = title_p.add_run("A CAPSTONE PROJECT REPORT\n\n")
    run.font.size = Pt(14)
    run.font.bold = True

    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author_p.add_run("SUBMITTED BY:\n")
    run.font.bold = True
    run = author_p.add_run("KUSHAL DEVRAJ (BU22CSEN0102284)\n\n")
    run = author_p.add_run("UNDER THE GUIDANCE OF:\n")
    run.font.bold = True
    run = author_p.add_run("SANGAMESHWAR\nAssistant Professor\n\n")
    run = author_p.add_run("DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING\nGITAM SCHOOL OF TECHNOLOGY\n2026\n")
    
    doc.add_page_break()

    # 2. ABSTRACT
    doc.add_heading('ABSTRACT', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(
        "Modern medical reports often present a significant barrier to health literacy due to high-density clinical jargon. "
        "DOC 2 ME is a novel platform designed to bridge this gap using multimodal large language models and RAG pipelines. "
        "The system incorporates automated PII scrubbing for privacy and Tesseract-based OCR for digitization. "
        "Our results show a 98.8% PII redaction accuracy and high semantic faithfulness in jargon simplification."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_page_break()

    # 3. CHAPTER 1: INTRODUCTION
    doc.add_heading('CHAPTER 1: INTRODUCTION', 1)
    p = doc.add_paragraph(
        "Health literacy is a critical determinant of patient outcomes. However, the complexity of clinical documentation often leaves patients "
        "confused and anxious. This project introduces 'Doc 2 Me', an AI-powered system designed to translate complex medical documents into plain language."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 4. CHAPTER 2: LITERATURE SURVEY
    doc.add_heading('CHAPTER 2: LITERATURE SURVEY', 1)
    p = doc.add_paragraph(
        "We surveyed 31 key research papers covering Transformers, Medical NLP, and de-identification techniques. "
        "While models like BioBERT excel in NER, generative models like T5 and Llama-3 provide superior abstractive simplification."
    )
    
    # 5. CHAPTER 3: METHODOLOGY
    doc.add_heading('CHAPTER 3: METHODOLOGY', 1)
    doc.add_heading('3.1 Data Acquisition', 2)
    
    # Table 1: MedDoc-1k
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parameter'
    hdr_cells[1].text = 'Value'
    data = [('Total Documents', '1,200'), ('Medical Categories', '8 (CBC, Radiology, etc.)'), ('PII Scrubbing', 'Regex + NER'), ('RAG Knowledge Base', '12,500+ Mappings')]
    for param, val in data:
        row_cells = table.add_row().cells
        row_cells[0].text = param
        row_cells[1].text = val
    
    doc.add_heading('3.2 Architecture', 2)
    if os.path.exists(os.path.join(img_dir, 'use_case_diagram.jpg')):
        doc.add_picture(os.path.join(img_dir, 'use_case_diagram.jpg'), width=Inches(5))
    doc.add_paragraph("Fig 3.1 Use Case Diagram Overview").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('3.3 Generative RAG Pipeline', 2)
    if os.path.exists(os.path.join(img_dir, 'fig3_4_rag_pipeline.jpg')):
        doc.add_picture(os.path.join(img_dir, 'fig3_4_rag_pipeline.jpg'), width=Inches(5))
    doc.add_paragraph("Fig 3.4 RAG Pipeline Workflow").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 6. CHAPTER 4: TESTING
    doc.add_heading('CHAPTER 4: TESTING', 1)
    p = doc.add_paragraph("Testing was performed across unit, integration, and performance layers. PII scrubbing achieved 98.8% accuracy.")

    # 7. CHAPTER 5: RESULTS
    doc.add_heading('CHAPTER 5: RESULTS', 1)
    if os.path.exists(os.path.join(img_dir, 'fig7_hematology.jpg')):
        doc.add_picture(os.path.join(img_dir, 'fig7_hematology.jpg'), width=Inches(4))
    doc.add_paragraph("Fig 5.1 Hematology Simplification Metrics").alignment = WD_ALIGN_PARAGRAPH.CENTER

    if os.path.exists(os.path.join(img_dir, 'fig10_query_distribution.jpg')):
        doc.add_picture(os.path.join(img_dir, 'fig10_query_distribution.jpg'), width=Inches(4))
    doc.add_paragraph("Fig 5.2 Patient Query Distribution").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 8. REFERENCES
    doc.add_heading('REFERENCES', 1)
    doc.add_paragraph("[1] World Health Organization, 2013.\n[2] Vaswani et al., 2017.\n[3] Devlin et al., 2018.\n...")

    # Set Global Font
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = 'Times New Roman'

    output_path = '/Users/kushaldevraj/Downloads/Medical_Jargon_Simplification-main/Doc2Me_Full_Report_Final.docx'
    doc.save(output_path)
    print(f"Final report saved at: {output_path}")

if __name__ == '__main__':
    build_final_document()
