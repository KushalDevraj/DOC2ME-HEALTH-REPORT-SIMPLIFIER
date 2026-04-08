from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_generic_report(category, filename_base, header_title, orig_text, results_data, impression_text):
    base_path = "/Users/kushaldevraj/Downloads/Medical_Jargon_Simplification-main/"
    
    # 1. Generate TEXT
    txt_path = f"{base_path}{filename_base}.txt"
    with open(txt_path, "w") as f:
        f.write(f"{header_title}\n{'-'*50}\n{orig_text}\n\nRESULTS:\n")
        for p, r, u, ref in results_data:
            f.write(f"{p}: {r} {u} (Ref: {ref})\n")
        f.write(f"\nIMPRESSION: {impression_text}\n")
    print(f"Created {txt_path}")

    # 2. Generate DOCX
    doc = Document()
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(f"CITY DIAGNOSTIC CENTER - {category.upper()}\n")
    run.font.size = Pt(16)
    run.font.bold = True
    doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading(header_title, 2)
    results_table = doc.add_table(rows=1, cols=4)
    results_table.style = 'Table Grid'
    hdr = results_table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "PARAMETER", "RESULT", "UNITS", "REF. RANGE"
    for p, r, u, ref in results_data:
        row = results_table.add_row().cells
        row[0].text, row[1].text, row[2].text, row[3].text = p, str(r), u, ref
    
    doc.add_paragraph("\n")
    imp = doc.add_paragraph()
    imp.add_run("IMPRESSION: ").font.bold = True
    imp.add_run(impression_text)
    
    docx_path = f"{base_path}{filename_base}.docx"
    doc.save(docx_path)
    print(f"Created {docx_path}")

    # 3. Generate JPEG
    width, height = 2480, 3508
    image = Image.new('RGB', (width, height), color='white')
    draw = ImageDraw.Draw(image)
    try:
        font_bold = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 80)
        font_reg = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 60)
        font_header = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 100)
    except:
        font_bold = font_reg = font_header = ImageFont.load_default()

    margin = 150
    y = 200
    draw.text((width/2, y), f"CITY DIAGNOSTIC CENTER - {category.upper()}", fill='black', font=font_header, anchor="mm")
    y += 150
    draw.line((margin, y, width-margin, y), fill='black', width=5)
    y += 150
    draw.text((margin, y), header_title, fill='black', font=font_header)
    y += 200
    
    for p, r, u, ref in results_data:
        draw.text((margin, y), f"{p}: {r} {u}", fill='black', font=font_reg)
        y += 100
    
    y += 150
    draw.text((margin, y), "IMPRESSION:", fill='black', font=font_bold)
    y += 100
    # Very basic wrapping
    words = impression_text.split()
    line = ""
    for word in words:
        if len(line + word) < 55: line += word + " "
        else:
            draw.text((margin + 50, y), line, fill='black', font=font_reg)
            y += 80
            line = word + " "
    draw.text((margin + 50, y), line, fill='black', font=font_reg)
    
    jpg_path = f"{base_path}{filename_base}.jpg"
    image.save(jpg_path, quality=95)
    print(f"Created {jpg_path}")

# --- DATA FOR REPORTS ---

# 1. RADIOLOGY
create_generic_report(
    "Radiology", "Sample_Radiology_Report", "CHEST X-RAY (PA VIEW)",
    "Patient presenting with persistent cough and tachycardia.",
    [("CTR", "0.62 *", "Ratio", "< 0.50"), ("Costophrenic Angles", "Blunted *", "", "Clear"), ("Lung Fields", "Bibasilar Opacity *", "", "Clear")],
    "Findings suggestive of Cardiomegaly and minor Bilateral Pleural Effusion. Clinical correlation with cardiac profile recommended."
)

# 2. BIOCHEMISTRY (LFT)
create_generic_report(
    "Biochemistry", "Sample_LFT_Report", "LIVER FUNCTION TEST (LFT)",
    "Patient presenting with abdominal pain and jaundice.",
    [("SGOT (AST)", "145 *", "U/L", "5 - 40"), ("SGPT (ALT)", "180 *", "U/L", "7 - 56"), ("Bilirubin Total", "3.2 *", "mg/dL", "0.1 - 1.2")],
    "Elevated transaminases and hyperbilirubinemia. Suggestive of acute hepatocellular injury/hepatitis."
)

# 3. CARDIOLOGY (ECG)
create_generic_report(
    "Cardiology", "Sample_ECG_Report", "12-LEAD ECG SUMMARY",
    "ECG obtained for chest discomfort and palpitations.",
    [("Heart Rate", "115 *", "BPM", "60 - 100"), ("PR Interval", "0.14", "sec", "0.12 - 0.20"), ("ST Segment", "Elevation *", "V1-V4", "Isoelectric")],
    "Sinus Tachycardia with ST-segment elevation in precordial leads. Possible acute anterior wall Myocardial Infarction. EMERGENCY correlation required."
)
