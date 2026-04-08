from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_generic_report(category, filename_base, header_title, orig_text, results_data, impression_text):
    base_path = "/Users/kushaldevraj/Downloads/Medical_Jargon_Simplification-main/"
    
    # 1. Generate TEXT
    txt_path = f"{base_path}{filename_base}.txt"
    with open(txt_path, "w") as f:
        f.write(f"{header_title}\n{'-'*50}\n{orig_text}\n\nRESULTS:\n")
        for p, r, u, ref in results_data:
            f.write(f"{p}: {r} {u} (Ref: {ref})\n")
        f.write(f"\nIMPRESSION: {impression_text}\n")

    # 2. Generate DOCX
    doc = Document()
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(f"GLOBAL HEALTH LABS - {category.upper()}\n")
    run.font.size = Pt(16)
    run.font.bold = True
    doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(header_title, 2)
    results_table = doc.add_table(rows=1, cols=4)
    results_table.style = 'Table Grid'
    hdr = results_table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "PARAMETER", "RESULT", "UNITS", "REF. RANGE"
    for p, r, u, ref in results_data:
        row = results_table.add_row().cells
        row[0].text, row[1].text, row[2].text, row[3].text = p, str(r), u, ref
    doc.add_paragraph("\n")
    imp = doc.add_paragraph()
    imp.add_run("IMPRESSION: ").font.bold = True
    imp.add_run(impression_text)
    doc.save(f"{base_path}{filename_base}.docx")

    # 3. Generate JPEG
    width, height = 2480, 3508
    image = Image.new('RGB', (width, height), color='white')
    draw = ImageDraw.Draw(image)
    try:
        font_bold = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 80)
        font_reg = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 60)
        font_header = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 100)
    except: font_bold = font_reg = font_header = ImageFont.load_default()
    margin, y = 150, 200
    draw.text((width/2, y), f"GLOBAL HEALTH LABS - {category.upper()}", fill='black', font=font_header, anchor="mm")
    y += 150
    draw.line((margin, y, width-margin, y), fill='black', width=5)
    y += 150
    draw.text((margin, y), header_title, fill='black', font=font_header)
    y += 200
    for p, r, u, ref in results_data:
        draw.text((margin, y), f"{p}: {r} {u}", fill='black', font=font_reg)
        y += 120
    y += 150
    draw.text((margin, y), "IMPRESSION:", fill='black', font=font_bold)
    y += 100
    words = impression_text.split()
    line = ""
    for word in words:
        if len(line + word) < 55: line += word + " "
        else:
            draw.text((margin + 50, y), line, fill='black', font=font_reg)
            y += 80
            line = word + " "
    draw.text((margin + 50, y), line, fill='black', font=font_reg)
    image.save(f"{base_path}{filename_base}.jpg", quality=95)

# 1. URINE ANALYSIS
create_generic_report(
    "Urine Analysis", "Sample_Urine_Report", "ROUTINE URINE ANALYSIS",
    "Clinical correlation with symptoms of dysuria.",
    [("Appearance", "Turbid *", "", "Clear"), ("Protein", "2+ *", "mg/dL", "Negative"), ("Pus Cells", "20-25 *", "/HPF", "0-5"), ("Bacteria", "Present *", "", "Absent")],
    "Findings suggestive of Urinary Tract Infection (UTI) and significant Proteinuria as evidenced by pyuria and bacteriuria."
)

# 2. LIPID PROFILE
create_generic_report(
    "Cardiology/Lipids", "Sample_Lipid_Report", "LIPID PROFILE (FASTING)",
    "Screening for cardiovascular risk factors.",
    [("Total Cholesterol", "265 *", "mg/dL", "< 200"), ("LDL (Bad)", "185 *", "mg/dL", "< 100"), ("HDL (Good)", "32 *", "mg/dL", "> 40"), ("Triglycerides", "240 *", "mg/dL", "< 150")],
    "Dyslipidemia with high LDL and low HDL levels. Increased risk for Atherosclerosis. Statins and lifestyle intervention recommended."
)

# 3. KFT (Kidney Function Test)
create_generic_report(
    "Renal Profile", "Sample_KFT_Report", "KIDNEY FUNCTION TEST (KFT)",
    "Patient presenting with persistent edema and hypertension.",
    [("Serum Creatinine", "2.8 *", "mg/dL", "0.7 - 1.3"), ("BUN", "52 *", "mg/dL", "7 - 20"), ("Uric Acid", "9.2 *", "mg/dL", "3.4 - 7.0")],
    "Significant Renal Impairment noted with Azotemia and Hyperuricemia. Correlation with GFR and Ultrasound KUB is required."
)
print("Bulk reports v2 generated.")
