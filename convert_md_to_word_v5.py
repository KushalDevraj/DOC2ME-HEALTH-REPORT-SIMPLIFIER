
import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "2"
    
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    run._r.append(fldChar4)

def set_style_font(doc, style_name, font_name='Times New Roman', font_size=12, bold=False, italic=False):
    try:
        style = doc.styles[style_name]
        font = style.font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = bold
        font.italic = italic
        
        # XML manipulation for persistence
        if hasattr(style, '_element') and style._element.rPr:
            style._element.rPr.rFonts.set(qn('w:ascii'), font_name)
            style._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
            style._element.rPr.rFonts.set(qn('w:cs'), font_name)
    except KeyError:
        pass 

def set_document_font(doc, font_name='Times New Roman'):
    """Set the default font for the document styles"""
    # Normal and Body Text: 12pt, Justified
    for style_id in ['Normal', 'Body Text']:
        set_style_font(doc, style_id, font_name, 12)
        doc.styles[style_id].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.styles[style_id].paragraph_format.line_spacing = 1.5
    
    # Headings: 16pt (Titles/Headings), 14pt (Subheadings)
    set_style_font(doc, 'Heading 1', font_name, 16, bold=True)
    doc.styles['Heading 1'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.styles['Heading 1'].paragraph_format.space_after = Pt(12)
    
    set_style_font(doc, 'Heading 2', font_name, 14, bold=True)
    doc.styles['Heading 2'].paragraph_format.space_before = Pt(12)
    doc.styles['Heading 2'].paragraph_format.space_after = Pt(6)

    set_style_font(doc, 'Heading 3', font_name, 12, bold=True)
    set_style_font(doc, 'Heading 4', font_name, 12, bold=True, italic=True)
        
    # Lists
    for style_id in ['List Bullet', 'List Number', 'List Paragraph']:
        set_style_font(doc, style_id, font_name, 12)
        doc.styles[style_id].paragraph_format.line_spacing = 1.5
    
    # Captions
    set_style_font(doc, 'Caption', font_name, 10, italic=True)

def add_formatted_text(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Handle italics inside non-bold parts
            italic_parts = re.split(r'(\*.*?\*)', part)
            for ipart in italic_parts:
                if ipart.startswith('*') and ipart.endswith('*') and len(ipart) > 2:
                    run = paragraph.add_run(ipart[1:-1])
                    run.italic = True
                else:
                    if ipart: 
                        run = paragraph.add_run(ipart)
    
    # Ensure all runs in this paragraph match the style
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

def parse_markdown(md_text, doc):
    lines = md_text.split('\n')
    in_code_block = False
    code_content = []
    
    # Simple table parser state
    in_table = False
    table_headers = []
    table_rows = []

    for line in lines:
        line = line.strip()
        
        # Explicit Page Break Handling
        if line == '{{PAGE_BREAK}}':
            doc.add_page_break()
            continue

        # Code block handling
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_content))
                run.font.name = 'Courier New' # Keep code monospaced
                run.font.size = Pt(10)
                p.style = 'No Spacing' 
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            continue
            
        if in_code_block:
            code_content.append(line)
            continue

        # Image Handling: ![Caption](path)
        img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
        if img_match:
            caption = img_match.group(1)
            img_path = img_match.group(2)
            
            # Resolve path relative to current directory if needed
            if not os.path.isabs(img_path):
                img_path = os.path.abspath(img_path)
                
            if os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(6.0))
                    last_paragraph = doc.paragraphs[-1] 
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    if caption:
                        caption_p = doc.add_paragraph(caption)
                        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_p.style = 'Caption'
                except Exception as e:
                    print(f"Error adding image {img_path}: {e}")
                    doc.add_paragraph(f"[Image: {caption} - Failed to load]")
            else:
                doc.add_paragraph(f"[Image: {caption} - File not found: {img_path}]")
            continue

        # Table handling
        if line.startswith('|') and line.endswith('|'):
            if '---' in line: 
                continue
            
            cells = [c.strip() for c in line.split('|') if c]
            if not in_table:
                in_table = True
                table_headers = cells
            else:
                table_rows.append(cells)
            continue
        elif in_table:
            # End of table, render it now
            table = doc.add_table(rows=1, cols=len(table_headers))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(table_headers):
                if i < len(hdr_cells):
                    hdr_cells[i].text = h
                    # Header formatting: Bold, Center
                    for p in hdr_cells[i].paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for r in p.runs:
                            r.bold = True
                            r.font.name = 'Times New Roman'
                            r.font.size = Pt(12)

            for row_data in table_rows:
                row = table.add_row().cells
                for i, cell_data in enumerate(row_data):
                    if i < len(row):
                        row[i].text = cell_data
                        for p in row[i].paragraphs:
                             p.alignment = WD_ALIGN_PARAGRAPH.CENTER # Center align table cells usually looks better for matrices
                             for r in p.runs:
                                r.font.name = 'Times New Roman'
                                r.font.size = Pt(11) # Slightly smaller for complex tables
            
            doc.add_paragraph() # Spacer
            in_table = False
            table_headers = []
            table_rows = []
        
        if not line:
            continue

        # Headers
        if line.startswith('# '):
            # If "Chapter" is in line, optional extra spacing or page break (handled by explicit tag now)
            h = doc.add_heading(line[2:], level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in h.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0,0,0)
                
        elif line.startswith('## '):
            h = doc.add_heading(line[3:], level=2)
            for run in h.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0,0,0)

        elif line.startswith('### '):
            h = doc.add_heading(line[4:], level=3)
            for run in h.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0,0,0)
        
        # Lists
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, line[2:])
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
            
        # Blockquotes
        elif line.startswith('> '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            p.paragraph_format.left_indent = Inches(0.5)

        # Standard Paragraph
        else:
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.5
            add_formatted_text(p, line)

    # Catch trailing table
    if in_table:
        table = doc.add_table(rows=1, cols=len(table_headers))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(table_headers):
            if i < len(hdr_cells):
                 hdr_cells[i].text = h
        for row_data in table_rows:
            row = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                if i < len(row):
                    row[i].text = cell_data

if __name__ == "__main__":
    try:
        with open('PROJECT_REPORT_FINAL.md', 'r', encoding='utf-8') as f:
            content = f.read()
            
        document = Document()
        set_document_font(document, 'Times New Roman')
        
        parse_markdown(content, document)
        
        # Add footer with page number
        section = document.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        add_page_number(run)
        
        output_file = 'Doc2Me_Health_Report_Expanded.docx'
        document.save(output_file)
        print(f"Successfully created {output_file}")
    except Exception as e:
        print(f"Error: {e}")
