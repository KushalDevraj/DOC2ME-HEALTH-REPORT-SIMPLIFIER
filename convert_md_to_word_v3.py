
import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_style_font(doc, style_name, font_name='Times New Roman', font_size=12, color=None):
    try:
        style = doc.styles[style_name]
        font = style.font
        font.name = font_name
        font.size = Pt(font_size)
        if color:
            font.color.rgb = color
        
        # XML manipulation for persistence
        if hasattr(style, '_element') and style._element.rPr:
            style._element.rPr.rFonts.set(qn('w:ascii'), font_name)
            style._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
            style._element.rPr.rFonts.set(qn('w:cs'), font_name)
    except KeyError:
        pass # Style might not exist

def set_document_font(doc, font_name='Times New Roman'):
    """Set the default font for the document styles"""
    # Normal text
    set_style_font(doc, 'Normal', font_name, 12)
    set_style_font(doc, 'Body Text', font_name, 12)
    
    # Headings
    for i in range(1, 5):
        set_style_font(doc, f'Heading {i}', font_name, 14 if i==1 else 13, RGBColor(0, 0, 0))
        
    # Lists
    set_style_font(doc, 'List Bullet', font_name, 12)
    set_style_font(doc, 'List Number', font_name, 12)
    set_style_font(doc, 'List Paragraph', font_name, 12)
    
    # Captions
    set_style_font(doc, 'Caption', font_name, 10)

def add_formatted_text(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        else:
            # Handle italics inside non-bold parts
            italic_parts = re.split(r'(\*.*?\*)', part)
            for ipart in italic_parts:
                if ipart.startswith('*') and ipart.endswith('*') and len(ipart) > 2:
                    run = paragraph.add_run(ipart[1:-1])
                    run.italic = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                else:
                    if ipart: # Avoid empty runs
                        run = paragraph.add_run(ipart)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)

def parse_markdown(md_text, doc):
    lines = md_text.split('\n')
    in_code_block = False
    code_content = []
    
    # Simple table parser state
    in_table = False
    table_headers = []
    table_rows = []

    for line in lines:
        line = line.strip()
        
        # Code block handling
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_content))
                run.font.name = 'Courier New' # Keep code monospaced
                run.font.size = Pt(10)
                p.style = 'No Spacing' 
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            continue
            
        if in_code_block:
            code_content.append(line)
            continue

        # Image Handling: ![Caption](path)
        img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
        if img_match:
            caption = img_match.group(1)
            img_path = img_match.group(2)
            
            # Resolve path relative to current directory if needed
            if not os.path.isabs(img_path):
                img_path = os.path.abspath(img_path)
                
            if os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(6.0))
                    last_paragraph = doc.paragraphs[-1] 
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    if caption:
                        caption_p = doc.add_paragraph(caption)
                        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # Manually style caption to avoid style issues
                        caption_run = caption_p.add_run() # Clear existing runs if any? No, add_paragraph(text) creates a run
                        # Actually add_paragraph(caption) creates a run with caption
                        caption_p.runs[0].font.name = 'Times New Roman'
                        caption_p.runs[0].font.size = Pt(10)
                        caption_p.runs[0].italic = True
                except Exception as e:
                    print(f"Error adding image {img_path}: {e}")
                    doc.add_paragraph(f"[Image: {caption} - Failed to load]")
            else:
                doc.add_paragraph(f"[Image: {caption} - File not found: {img_path}]")
            continue

        # Table handling (basic)
        if line.startswith('|') and line.endswith('|'):
            if '---' in line: # Separator line
                continue
            
            cells = [c.strip() for c in line.split('|') if c]
            if not in_table:
                in_table = True
                table_headers = cells
            else:
                table_rows.append(cells)
            continue
        elif in_table:
            # End of table, render it now
            table = doc.add_table(rows=1, cols=len(table_headers))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(table_headers):
                if i < len(hdr_cells):
                    hdr_cells[i].text = h
                    for paragraph in hdr_cells[i].paragraphs:
                        if paragraph.runs:
                            paragraph.runs[0].font.name = 'Times New Roman'
                            paragraph.runs[0].font.size = Pt(12)
                            paragraph.runs[0].bold = True
            
            for row_data in table_rows:
                row = table.add_row().cells
                for i, cell_data in enumerate(row_data):
                    if i < len(row):
                        row[i].text = cell_data
                        for paragraph in row[i].paragraphs:
                             if paragraph.runs:
                                paragraph.runs[0].font.name = 'Times New Roman'
                                paragraph.runs[0].font.size = Pt(12)
            
            doc.add_paragraph() # Spacer
            in_table = False
            table_headers = []
            table_rows = []
        
        if not line:
            continue

        # Headers
        if line.startswith('# '):
            h = doc.add_heading(line[2:], level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in h.runs:
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0,0,0)
                
        elif line.startswith('## '):
            h = doc.add_heading(line[3:], level=2)
            for run in h.runs:
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0,0,0)

        elif line.startswith('### '):
            h = doc.add_heading(line[4:], level=3)
            for run in h.runs:
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0,0,0)

        elif line.startswith('#### '):
            h = doc.add_heading(line[5:], level=4)
            for run in h.runs:
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0,0,0)
        
        # Lists
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, line[2:])
        elif re.match(r'^\d+\.\s', line):
            # Numbered list
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
            
        # Blockquotes/Notes
        elif line.startswith('> '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            p.paragraph_format.left_indent = Inches(0.5)

        # Standard Paragraph
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, line)

    # Catch trailing table if file ends with one
    if in_table:
        table = doc.add_table(rows=1, cols=len(table_headers))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(table_headers):
            if i < len(hdr_cells):
                 hdr_cells[i].text = h
        for row_data in table_rows:
            row = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                if i < len(row):
                    row[i].text = cell_data

if __name__ == "__main__":
    try:
        with open('PROJECT_REPORT_FINAL.md', 'r', encoding='utf-8') as f:
            content = f.read()
            
        document = Document()
        set_document_font(document, 'Times New Roman')
        
        # Add Title with explicit font styling
        title = document.add_heading('MediClare - Medical Jargon Simplification Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        parse_markdown(content, document)
        
        output_file = 'MediClare_Project_Report_v3.docx'
        document.save(output_file)
        print(f"Successfully created {output_file}")
    except Exception as e:
        print(f"Error: {e}")
