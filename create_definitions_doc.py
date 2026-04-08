import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# Set default font for the document to Times New Roman, 12pt
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Fix for Mac MS Word font rendering 
if hasattr(style, '_element') and style._element.rPr:
    style._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    style._element.rPr.rFonts.set(qn('w:cs'), 'Times New Roman')

# Set standard research paper paragraph formatting (Justified, 1.5 line spacing)
paragraph_format = doc.styles['Normal'].paragraph_format
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
paragraph_format.line_spacing = 1.5
paragraph_format.space_after = Pt(12)

# The generated definitions
definitions = [
    ("Medical Text Simplification", "Medical text simplification is an NLP process that translates complex clinical terminology—a set of specialized medical vocabulary and dense syntax—into plain, understandable English. A system performs medical text simplification if it successfully reduces syntactic complexity while strictly preserving the underlying semantic meaning. This process is driven by the need to democratize access to health records and improve patient health literacy."),
    ("Retrieval-Augmented Generation (RAG)", "Retrieval-Augmented Generation is an architectural framework that grounds language models by querying an external verified database—a high-dimensional vector index of clinical knowledge—before generating a response. An AI conversational agent utilizes RAG if it fetches factual context to construct its answers rather than relying solely on its internal parametric memory. Unlike standard unstructured generative models, this significantly reduces the risk of the model hallucinating dangerous or incorrect medical advice."),
    ("Text-To-Text Transfer Transformer (T5)", "A Text-To-Text Transfer Transformer is a sequence-to-sequence machine learning model that casts all natural language processing tasks—including translation, summarization, and simplification—as text generation problems. A machine learning model operates as a fine-tuned T5 if it has been specialized to correctly map complex medical input strings directly to their layperson output equivalents. This architecture is based on Google's unified framework for NLP tasks."),
    ("Personally Identifiable Information (PII) Scrubbing", "PII Scrubbing is a privacy-preserving technique that aggressively detects and redacts sensitive patient data—such as names, ages, and contact information—before any AI inference occurs. A system implements PII scrubbing if it successfully identifies these patterns and replaces them with generic placeholder tokens. Because this step occurs prior to vector embedding, it ensures no sensitive information is ever memorized by the underlying neural networks."),
    ("Local-First Architecture", "A local-first architecture is a system design paradigm where all computational tasks and large language model inferences occur entirely on the user's host machine—bypassing the need for external public APIs. An application is local-first if no sensitive medical data is transmitted over the internet during its usage. This architectural decision strictly guarantees patient confidentiality and solves the primary privacy hurdle facing cloud-based healthcare AI.")
]

# Add a title
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run("Core Concepts and Definitions")
title_run.bold = True
title_run.font.size = Pt(14)
title_para.paragraph_format.space_after = Pt(24)

# Create the paragraphs
for title, text in definitions:
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Add title in bold
    run_title = p.add_run(f"{title}:")
    run_title.bold = True
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(12)
    
    # Add Space and text
    run_text = p.add_run(f" {text}")
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(12)

# Save the document
output_file = 'Doc2Me_Definitions.docx'
doc.save(output_file)
print(f"Successfully created {output_file}")
